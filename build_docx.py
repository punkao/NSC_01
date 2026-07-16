#!/usr/bin/env python3
"""สร้างรายงาน .docx จาก REPORT_SECTION5.md — ทำเป็นสคริปต์เพื่อให้ generate ใหม่ได้ทุกครั้งที่แก้ .md

ทำไมต้องมี: .docx ฉบับก่อนสร้างด้วยมือ พอแก้ตัวเลขใน .md แล้ว .docx ไม่ตาม
-> ตอนจะส่งเลยกลายเป็นส่งเวอร์ชันที่ตัวเลขผิดทั้งหมด. สคริปต์นี้ตัดปัญหานั้นถาวร

ฟอร์แมตตามข้อกำหนด: A4 · ขอบ 1 นิ้ว · TH Sarabun New 16pt · มีเลขหน้า
(ฉบับก่อนหน้า footer ว่าง = ไม่มีเลขหน้า ซึ่งผิดข้อกำหนด — ฉบับนี้ใส่ให้แล้ว)

Run: python build_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SRC = Path("REPORT_SECTION5.md")
OUT = Path("รายงาน_บทที่5_Line-Guard.docx")
FONT = "TH Sarabun New"
BASE_PT = 16

COVER = [
    ("รายงานฉบับสมบูรณ์", 24, True),
    ("", 16, False),
    ("Line Guard", 30, True),
    ("ระบบตรวจจับเหตุการณ์ความปลอดภัยในสายการผลิต", 22, True),
    ("ด้วยแบบจำลองภาษาและภาพ (Vision-Language Model)", 22, True),
    ("", 16, False),
    ("Line Guard: A Vision-Language Model Approach to", 18, False),
    ("Industrial Safety Event Detection from CCTV Video", 18, False),
    ("", 16, False),
    ("บทที่ 5  รายละเอียดของการพัฒนา", 20, True),
    ("", 16, False),
    ("", 16, False),
    ("จัดทำโดย", 18, True),
    ("[ระบุชื่อ - นามสกุล ผู้พัฒนา]", 18, False),
    ("[ระบุชื่อทีม / รหัสนักศึกษา]", 18, False),
    ("", 16, False),
    ("อาจารย์ที่ปรึกษา", 18, True),
    ("[ระบุชื่ออาจารย์ที่ปรึกษา]", 18, False),
    ("", 16, False),
    ("[ระบุชื่อสถาบันการศึกษา]", 18, False),
    ("[ระบุชื่อโครงการ / การแข่งขัน]  ปีการศึกษา [ระบุปี]", 18, False),
]


def set_font(run, size: int = BASE_PT, bold: bool = False, mono: bool = False) -> None:
    """ตั้งฟอนต์ให้ครบทั้ง ascii/hAnsi/cs — ภาษาไทยใช้ cs ถ้าไม่ตั้งจะเพี้ยน"""
    name = "Courier New" if mono else FONT
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    for tag in ("w:rFonts",):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            el.set(qn(attr), name)
    sz = OxmlElement("w:szCs")            # ขนาดฟอนต์ฝั่ง complex script (ไทย)
    sz.set(qn("w:val"), str(size * 2))
    rpr.append(sz)
    if bold:
        b = OxmlElement("w:bCs")
        rpr.append(b)


def add_runs(p, text: str, size: int = BASE_PT) -> None:
    """แปลง **ตัวหนา** และ `โค้ด` เป็น run — markdown inline แบบง่าย"""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_font(p.add_run(part[2:-2]), size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            set_font(p.add_run(part[1:-1]), size - 2, mono=True)
        else:
            set_font(p.add_run(part), size)


def add_page_number(section) -> None:
    """ใส่เลขหน้าใน footer (ข้อกำหนดบังคับ — ฉบับก่อนหน้าไม่มี)"""
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr, kind in (("begin", "w:fldChar"), ("PAGE", "w:instrText"), ("end", "w:fldChar")):
        el = OxmlElement(kind)
        if kind == "w:fldChar":
            el.set(qn("w:fldCharType"), instr)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        run._r.append(el)
    set_font(run, 14)


def md_table(doc, rows: list[str]) -> None:
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], [r for r in cells[2:]]     # ข้าม row คั่น |---|
    t = doc.add_table(rows=1 + len(body), cols=len(header))
    t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ""
        add_runs(cell.paragraphs[0], h or " ", BASE_PT - 2)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(body):
        for j, v in enumerate(row[:len(header)]):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            add_runs(cell.paragraphs[0], v or " ", BASE_PT - 2)


def main() -> None:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(BASE_PT)
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)      # A4
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))

    # ---- หน้าปก ----
    for text, size, bold in COVER:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), size, bold)
    doc.add_page_break()
    add_page_number(sec)

    # ---- เนื้อหาจาก markdown ----
    lines = SRC.read_text(encoding="utf-8").split("\n")
    i, in_code, buf_tbl = 0, False, []
    while i < len(lines):
        ln = lines[i]

        if ln.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            set_font(p.add_run(ln), BASE_PT - 3, mono=True)
            last_blank = False
            i += 1
            continue

        # ตาราง — เก็บทั้งบล็อกก่อนค่อยสร้าง
        if ln.strip().startswith("|"):
            buf_tbl.append(ln)
            if i + 1 >= len(lines) or not lines[i + 1].strip().startswith("|"):
                md_table(doc, buf_tbl)
                buf_tbl = []
                last_blank = False
            i += 1
            continue

        m_img = re.match(r"!\[[^\]]*\]\(([^)]+)\)", ln.strip())
        if m_img:
            path = Path(m_img.group(1))
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.3))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_blank = False
            i += 1
            continue

        s = ln.strip()
        if s.startswith("# "):
            p = doc.add_paragraph()
            set_font(p.add_run(s[2:]), 20, bold=True)
        elif s.startswith("## "):
            if not last_blank:
                doc.add_paragraph()
            p = doc.add_paragraph()
            set_font(p.add_run(s[3:]), 18, bold=True)
        elif s.startswith("### "):
            p = doc.add_paragraph()
            set_font(p.add_run(s[4:]), 17, bold=True)
        elif s.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            add_runs(p, s[2:], BASE_PT - 1)
            for r in p.runs:
                r.italic = True
        elif re.match(r"^[-*] ", s):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:])
        elif re.match(r"^\d+\. ", s):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", s))
        elif s == "---":
            i += 1                      # เส้นคั่นไม่สร้างอะไร และต้องไม่รีเซ็ต last_blank
            continue
        elif s:
            p = doc.add_paragraph()
            add_runs(p, s)
        else:
            if not last_blank:
                doc.add_paragraph()
            last_blank = True
            i += 1
            continue
        last_blank = False
        i += 1

    doc.save(OUT)
    print(f"สร้าง {OUT} แล้ว  ({len(doc.paragraphs)} ย่อหน้า, {len(doc.tables)} ตาราง, {len(doc.inline_shapes)} รูป)")


if __name__ == "__main__":
    main()
